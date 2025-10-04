import os
from docx import Document
from docx.shared import Inches
import pdfplumber
from typing import Dict, Any

class FileHandler:
    def __init__(self):
        pass
    
    def save_resume(self, resume_data: Dict[str, Any], output_path: str, original_path: str):
        """Save optimized resume in the same format as original"""
        if original_path.lower().endswith('.docx'):
            self._save_as_docx(resume_data, output_path)
        elif original_path.lower().endswith('.pdf'):
            self._save_as_pdf(resume_data, output_path)
        else:
            self._save_as_txt(resume_data, output_path)
    
    def _save_as_docx(self, resume_data: Dict[str, Any], output_path: str):
        """Save resume as DOCX"""
        doc = Document()
        
        # Add title
        title = doc.add_heading('Optimized Resume', 0)
        
        # Add personal info
        if resume_data.get('personal_info'):
            personal = resume_data['personal_info']
            doc.add_paragraph(f"Name: {personal.get('name', '')}")
            doc.add_paragraph(f"Email: {personal.get('email', '')}")
            doc.add_paragraph(f"Phone: {personal.get('phone', '')}")
            doc.add_paragraph()
        
        # Add summary
        if resume_data.get('summary'):
            doc.add_heading('Summary', level=1)
            doc.add_paragraph(resume_data['summary'])
            doc.add_paragraph()
        
        # Add skills
        if resume_data.get('skills'):
            doc.add_heading('Skills', level=1)
            skills_text = ', '.join(resume_data['skills'])
            doc.add_paragraph(skills_text)
            doc.add_paragraph()
        
        # Add experience
        if resume_data.get('experience'):
            doc.add_heading('Experience', level=1)
            for job in resume_data['experience']:
                # Job title and company
                job_title = job.get('title', '')
                company = job.get('company', '')
                dates = ' | '.join(job.get('dates', []))
                
                experience_line = f"{job_title}"
                if company:
                    experience_line += f" at {company}"
                if dates:
                    experience_line += f" ({dates})"
                
                doc.add_paragraph(experience_line, style='Heading 2')
                
                # Job description
                for bullet in job.get('description', []):
                    p = doc.add_paragraph(bullet, style='List Bullet')
        
        # Add education
        if resume_data.get('education'):
            doc.add_heading('Education', level=1)
            for edu in resume_data['education']:
                degree = edu.get('degree', '')
                institution = edu.get('institution', '')
                dates = ' | '.join(edu.get('dates', []))
                
                education_line = f"{degree}"
                if institution:
                    education_line += f" - {institution}"
                if dates:
                    education_line += f" ({dates})"
                
                doc.add_paragraph(education_line, style='Heading 2')
                
                for detail in edu.get('details', []):
                    doc.add_paragraph(detail, style='List Bullet')
        
        doc.save(output_path)
    
    def _save_as_pdf(self, resume_data: Dict[str, Any], output_path: str):
        """Save resume as PDF (simplified - in production, use reportlab or similar)"""
        # For now, save as text since PDF generation requires additional libraries
        # In a real implementation, you would use reportlab or weasyprint
        self._save_as_txt(resume_data, output_path.replace('.pdf', '.txt'))
    
    def _save_as_txt(self, resume_data: Dict[str, Any], output_path: str):
        """Save resume as text"""
        with open(output_path, 'w', encoding='utf-8') as f:
            # Personal info
            if resume_data.get('personal_info'):
                personal = resume_data['personal_info']
                f.write(f"Name: {personal.get('name', '')}\n")
                f.write(f"Email: {personal.get('email', '')}\n")
                f.write(f"Phone: {personal.get('phone', '')}\n\n")
            
            # Summary
            if resume_data.get('summary'):
                f.write("SUMMARY\n")
                f.write("=" * 50 + "\n")
                f.write(f"{resume_data['summary']}\n\n")
            
            # Skills
            if resume_data.get('skills'):
                f.write("SKILLS\n")
                f.write("=" * 50 + "\n")
                f.write(', '.join(resume_data['skills']) + "\n\n")
            
            # Experience
            if resume_data.get('experience'):
                f.write("EXPERIENCE\n")
                f.write("=" * 50 + "\n")
                for job in resume_data['experience']:
                    job_title = job.get('title', '')
                    company = job.get('company', '')
                    dates = ' | '.join(job.get('dates', []))
                    
                    f.write(f"{job_title}")
                    if company:
                        f.write(f" at {company}")
                    if dates:
                        f.write(f" ({dates})")
                    f.write("\n")
                    
                    for bullet in job.get('description', []):
                        f.write(f"  • {bullet}\n")
                    f.write("\n")
            
            # Education
            if resume_data.get('education'):
                f.write("EDUCATION\n")
                f.write("=" * 50 + "\n")
                for edu in resume_data['education']:
                    degree = edu.get('degree', '')
                    institution = edu.get('institution', '')
                    dates = ' | '.join(edu.get('dates', []))
                    
                    f.write(f"{degree}")
                    if institution:
                        f.write(f" - {institution}")
                    if dates:
                        f.write(f" ({dates})")
                    f.write("\n")
                    
                    for detail in edu.get('details', []):
                        f.write(f"  • {detail}\n")
                    f.write("\n")
